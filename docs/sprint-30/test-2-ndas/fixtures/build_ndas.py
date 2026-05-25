#!/usr/bin/env python3
"""Generate 10 NDA DOCX fixtures for Sprint 30 Test 2.

Each NDA varies along: mutual/bilateral, term length, definition of
Confidential Information, residuals carve-out, governing law,
counterparty industry. The variation is enough that a per-NDA review
should produce distinct GREEN/YELLOW/RED verdicts.

Output: ./staged/nda-<counterparty>.docx
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches


HERE = Path(__file__).resolve().parent
OUT = HERE / "staged"
OUT.mkdir(exist_ok=True)

OUR_NAME = "Stanford Industrial Supply Co."


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.size = Pt(13 if level == 1 else 12)


def render(filename, body):
    doc = Document()
    for para in body:
        if isinstance(para, tuple):
            kind = para[0]
            if kind == "h1":
                add_heading(doc, para[1], 1)
            elif kind == "h2":
                add_heading(doc, para[1], 2)
            elif kind == "bold":
                add_para(doc, para[1], bold=True)
            elif kind == "page":
                doc.add_page_break()
        else:
            add_para(doc, para)
    doc.save(OUT / filename)


def std_recitals(party_a, party_b, purpose):
    return [
        ("bold", f"DATED: ___ 2026"),
        "",
        ("bold", "PARTIES"),
        f"(1) {party_a}",
        f"(2) {party_b}",
        "",
        ("bold", "BACKGROUND"),
        f"The parties wish to discuss a potential business arrangement relating to {purpose} (the \"Purpose\"). In connection with those discussions each party may disclose Confidential Information to the other.",
    ]


def std_governing(law):
    return [
        ("h2", "Governing Law and Jurisdiction"),
        f"This Agreement is governed by {law} and the parties submit to the exclusive jurisdiction of the courts of that jurisdiction.",
    ]


def std_misc():
    return [
        ("h2", "Miscellaneous"),
        "No Party may assign without the other's prior written consent. This Agreement may be executed in counterparts including by electronic signature. The Agreement contains the entire understanding between the parties in relation to its subject matter and supersedes any prior arrangement or representation.",
    ]


# --------------------------------------------------------------------
# 1. ACME TOOLING — mutual, balanced, clean (GREEN expected)
# --------------------------------------------------------------------
render(
    "nda-acme-tooling.docx",
    [
        ("h1", "MUTUAL CONFIDENTIALITY AGREEMENT"),
        ("h2", "Parties and Purpose"),
        *std_recitals(
            OUR_NAME,
            "ACME Tooling Solutions Limited (registered in England and Wales no. 04123876)",
            "ACME's potential supply of specialist cutting tools and abrasives to Stanford"
        ),
        ("h2", "1. Definition of Confidential Information"),
        "\"Confidential Information\" means any information disclosed by one party (\"Discloser\") to the other (\"Recipient\") in connection with the Purpose that is marked confidential, identified as confidential at the time of disclosure, or that a reasonable person would understand to be confidential given its nature and the circumstances of disclosure.",
        ("h2", "2. Carve-outs"),
        "Confidential Information does not include information that (a) is or becomes publicly available without breach of this Agreement; (b) was already known to the Recipient without obligation of confidence; (c) is received from a third party without restriction; (d) is independently developed without reference to the Discloser's information; or (e) is required to be disclosed by law or court order, provided that the Recipient gives prompt notice and reasonable cooperation.",
        ("h2", "3. Obligations"),
        "The Recipient shall (a) use Confidential Information solely for the Purpose; (b) protect it with the same degree of care as it uses for its own confidential information, but not less than reasonable care; (c) limit access to those of its personnel and professional advisers who need to know for the Purpose, and who are bound by equivalent obligations of confidence.",
        ("h2", "4. Term"),
        "This Agreement commences on the date above and continues for two (2) years. Obligations of confidence in respect of Confidential Information disclosed during the term survive for a further three (3) years from the date of disclosure.",
        ("h2", "5. Residuals"),
        "Nothing in this Agreement restricts the Recipient's use of residual information retained in the unaided memory of its personnel after the term, provided no Confidential Information is intentionally retained.",
        ("h2", "6. No Licence"),
        "No licence or other right is granted under any intellectual property of the Discloser by virtue of this Agreement.",
        ("h2", "7. Return or Destruction"),
        "On request or termination, the Recipient shall return or destroy all Confidential Information in its possession (other than back-up copies retained in the ordinary course) and certify destruction on request.",
        *std_governing("the laws of England and Wales"),
        *std_misc(),
    ],
)


# --------------------------------------------------------------------
# 2. BETACORP LOGISTICS — bilateral, slight buyer-favourable (YELLOW)
# --------------------------------------------------------------------
render(
    "nda-betacorp-logistics.docx",
    [
        ("h1", "NON-DISCLOSURE AGREEMENT"),
        *std_recitals(
            "BetaCorp Logistics PLC (\"Discloser\")",
            f"{OUR_NAME} (\"Recipient\")",
            "BetaCorp's evaluation of Stanford as a potential supplier of MRO consumables to its UK distribution network"
        ),
        ("h2", "1. Definitions"),
        "\"Confidential Information\" means all information disclosed by Discloser to Recipient, in any form, in connection with the Purpose, whether or not marked or identified as confidential.",
        ("h2", "2. Obligations"),
        "Recipient shall (a) use Confidential Information solely for the Purpose; (b) protect it with the same degree of care it uses for its own most sensitive information; (c) restrict access to those personnel and advisers who have a need to know.",
        ("h2", "3. Carve-outs"),
        "Confidential Information does not include information that is or becomes publicly available without breach, was already known to Recipient without obligation of confidence, is received from a third party without restriction, or is independently developed.",
        ("h2", "4. Term"),
        "This Agreement commences on the date above and continues for five (5) years. The Recipient's obligations of confidence survive for a further seven (7) years after termination or expiry.",
        ("h2", "5. No Disclosure of Discussions"),
        "Recipient shall not disclose to any third party the existence or substance of any discussions between the parties without Discloser's prior written consent.",
        ("h2", "6. Return"),
        "On request, Recipient shall return or destroy all Confidential Information at Discloser's election. The Recipient shall certify destruction within five (5) Business Days.",
        ("h2", "7. Remedies"),
        "Recipient acknowledges that damages may not be an adequate remedy for breach and that Discloser is entitled to seek injunctive relief.",
        *std_governing("the laws of England and Wales"),
        *std_misc(),
    ],
)


# --------------------------------------------------------------------
# 3. CYPRESS SAAS — mutual; indefinite term (RED)
# --------------------------------------------------------------------
render(
    "nda-cypress-saas.docx",
    [
        ("h1", "MUTUAL NON-DISCLOSURE AGREEMENT"),
        *std_recitals(
            OUR_NAME,
            "Cypress Systems, Inc. (a Delaware corporation)",
            "Cypress's potential procurement of inventory management SaaS by Stanford"
        ),
        ("h2", "1. Confidential Information"),
        "\"Confidential Information\" means all non-public information of either party, including without limitation any business, financial, technical, customer, supplier, pricing, strategic, personnel, regulatory, or operational information, in any form, whether disclosed before or after the date of this Agreement.",
        ("h2", "2. Obligations"),
        "Each party shall use the other party's Confidential Information solely for the Purpose, shall not disclose it to any third party other than personnel and advisers who need to know for the Purpose, and shall protect it with at least the same degree of care as it uses for its own most sensitive information.",
        ("h2", "3. Term and Survival"),
        "This Agreement is **perpetual** and continues in force indefinitely. There is no fixed expiry. The obligations of confidence survive indefinitely after the cessation of any discussions between the parties.",
        ("h2", "4. Residuals"),
        "There is no carve-out for residual information retained in the unaided memory of personnel. All Confidential Information must be returned or destroyed on request.",
        ("h2", "5. No Reverse Engineering"),
        "Recipient shall not reverse engineer, decompile, or attempt to derive the source code or underlying algorithms of any software or process embodied in the Confidential Information.",
        ("h2", "6. Equitable Relief"),
        "Each party acknowledges that monetary damages are an inadequate remedy and the disclosing party shall be entitled to injunctive relief without bond or proof of irreparable harm.",
        *std_governing("the laws of the State of California, USA"),
        *std_misc(),
    ],
)


# --------------------------------------------------------------------
# 4. DELPHI COMPONENTS — bilateral, clean (GREEN)
# --------------------------------------------------------------------
render(
    "nda-delphi-components.docx",
    [
        ("h1", "CONFIDENTIALITY AGREEMENT"),
        *std_recitals(
            "Delphi Components GmbH (registered in Germany, HRB 234567)",
            f"{OUR_NAME}",
            "potential supply by Delphi of electrical components into Stanford's EU distribution network"
        ),
        ("h2", "1. Definition"),
        "\"Confidential Information\" means information that the disclosing party identifies in writing or orally at the time of disclosure as confidential.",
        ("h2", "2. Use Restriction"),
        "The receiving party shall use Confidential Information only for the purpose of evaluating the potential supply arrangement and shall not use it for any commercial purpose beyond that evaluation.",
        ("h2", "3. Carve-outs"),
        "The obligations in clause 2 do not apply to information that (a) is in the public domain through no fault of the receiving party; (b) was already known to the receiving party at the time of disclosure; (c) is received from a third party without restriction; or (d) is required to be disclosed by law, regulation, or order of a competent authority.",
        ("h2", "4. Term"),
        "This Agreement is valid for three (3) years from the date of execution. Obligations of confidence survive for a further two (2) years from the end of the term.",
        ("h2", "5. Return of Documents"),
        "On expiry or termination, the receiving party shall return or destroy all Confidential Information and certify destruction within ten (10) Business Days of request.",
        *std_governing("German law"),
        *std_misc(),
    ],
)


# --------------------------------------------------------------------
# 5. EMBERLAKE COBIDDER — mutual; broad CI definition (YELLOW)
# --------------------------------------------------------------------
render(
    "nda-emberlake-cobidder.docx",
    [
        ("h1", "MUTUAL NON-DISCLOSURE AGREEMENT"),
        *std_recitals(
            "Emberlake Holdings Limited",
            f"{OUR_NAME}",
            "the parties' joint evaluation of a co-bidding opportunity for a public-sector framework agreement"
        ),
        ("h2", "1. Confidential Information"),
        "\"Confidential Information\" means any and all information, technical or non-technical, in any form or medium, of any kind, that is disclosed or made available by one party to the other in any context or manner, whether or not marked confidential and whether or not the disclosing party considers it confidential at the time, **including information that is publicly available but compiled, organised, or contextualised by the disclosing party**. This definition extends to information that the receiving party could reasonably be expected to know is sensitive to the disclosing party regardless of how it was acquired.",
        ("h2", "2. Obligations"),
        "Each party shall use Confidential Information only for the Purpose and shall not disclose it to any third party other than personnel and advisers with a need to know.",
        ("h2", "3. Limited Carve-outs"),
        "The only exception to clause 2 is information that the receiving party can demonstrate by contemporaneous written records was already in its possession without obligation of confidence prior to disclosure.",
        ("h2", "4. Term"),
        "This Agreement commences on signature and continues for the duration of any discussions between the parties relating to the Purpose plus three (3) years thereafter.",
        ("h2", "5. Non-Solicitation"),
        "During the term and for twelve (12) months thereafter, neither party shall directly or indirectly solicit for employment any employee of the other party with whom it had contact in connection with the Purpose.",
        ("h2", "6. No Public Statements"),
        "Neither party shall make any public statement about the existence or substance of the discussions without the other party's prior written consent.",
        *std_governing("the laws of England and Wales"),
        *std_misc(),
    ],
)


# --------------------------------------------------------------------
# 6. FOLIUM RESEARCH — mutual, balanced (GREEN)
# --------------------------------------------------------------------
render(
    "nda-folium-research.docx",
    [
        ("h1", "MUTUAL CONFIDENTIALITY UNDERTAKING"),
        *std_recitals(
            f"{OUR_NAME}",
            "Folium Research Institute (a not-for-profit company limited by guarantee)",
            "potential collaboration on a sustainable-packaging research project funded by Innovate UK"
        ),
        ("h2", "1. Definitions"),
        "\"Confidential Information\" means information disclosed by one party to the other that is identified as confidential at the time of disclosure, whether marked in writing or stated orally and confirmed in writing within 14 days.",
        ("h2", "2. Use and Disclosure"),
        "Each party shall (a) use the other's Confidential Information only for the Purpose, (b) limit access to its personnel and advisers with a need to know, and (c) protect it with reasonable care.",
        ("h2", "3. Exclusions"),
        "Confidential Information does not include information that (a) is or becomes publicly available without breach of this undertaking; (b) was known to the receiving party prior to disclosure; (c) is received from a third party without breach of any obligation of confidence; (d) is independently developed without reference to the disclosing party's information; or (e) must be disclosed under law or regulation (with prompt notice to the disclosing party).",
        ("h2", "4. Term"),
        "This undertaking shall remain in force for two (2) years from the date of last signature. Obligations of confidence survive for a further three (3) years from the date of disclosure.",
        ("h2", "5. Publication"),
        "Where the Purpose involves academic or industry publication, each party shall give the other thirty (30) days to review and comment on any proposed publication that may contain the other's Confidential Information.",
        ("h2", "6. Return or Destruction"),
        "On request or on completion of the Purpose, each party shall return or destroy the other's Confidential Information except for one copy that may be retained for legal or regulatory record-keeping.",
        *std_governing("the laws of England and Wales"),
        *std_misc(),
    ],
)


# --------------------------------------------------------------------
# 7. GREENLINE ERP — bilateral; no residuals carve-out, US law (RED)
# --------------------------------------------------------------------
render(
    "nda-greenline-erp.docx",
    [
        ("h1", "NON-DISCLOSURE AGREEMENT"),
        *std_recitals(
            "GreenLine ERP Systems LLC (a Texas limited liability company)",
            f"{OUR_NAME}",
            "GreenLine's potential supply of ERP integration services into Stanford's operations"
        ),
        ("h2", "1. Confidential Information"),
        "\"Confidential Information\" means any and all non-public information disclosed by GreenLine to Stanford in any form, including without limitation source code, algorithms, system architectures, customer lists, pricing methodologies, financial data, supplier identities, and any derivatives or compilations of any of the foregoing, regardless of whether marked confidential.",
        ("h2", "2. Use"),
        "Stanford shall use Confidential Information solely for the evaluation of GreenLine's proposed services. Stanford shall not use Confidential Information for the development of any competing product or service, whether internally or with any third party.",
        ("h2", "3. Survival"),
        "Stanford's obligations under this Agreement shall survive in perpetuity. There is no termination date and the obligations do not lapse.",
        ("h2", "4. No Residuals"),
        "There is no carve-out for residual information retained in the unaided memory of Stanford's personnel. Any information acquired through this Agreement remains subject to the obligations of confidence indefinitely whether or not retained in tangible form.",
        ("h2", "5. Non-Compete"),
        "For a period of two (2) years following any termination of discussions, Stanford shall not engage with any third-party supplier of ERP services that competes with GreenLine in respect of any line of business in which GreenLine has disclosed Confidential Information.",
        ("h2", "6. Liquidated Damages"),
        "The parties agree that any breach of this Agreement by Stanford shall give rise to liquidated damages of US$500,000 per incident, which the parties agree is a reasonable estimate of the loss the disclosing party would suffer.",
        ("h2", "7. Equitable Relief"),
        "GreenLine shall be entitled to injunctive relief in any court of competent jurisdiction without bond or proof of irreparable harm.",
        *std_governing("the laws of the State of Texas, USA, without regard to its conflict-of-laws principles"),
        *std_misc(),
    ],
)


# --------------------------------------------------------------------
# 8. HARBORWAVE CLOUD — mutual, clean (GREEN)
# --------------------------------------------------------------------
render(
    "nda-harborwave-cloud.docx",
    [
        ("h1", "MUTUAL NON-DISCLOSURE AGREEMENT"),
        *std_recitals(
            f"{OUR_NAME}",
            "Harborwave Cloud Services BV (registered in the Netherlands, KvK 12345678)",
            "evaluation of cloud-hosted document management services for Stanford's contract repository"
        ),
        ("h2", "1. Confidential Information"),
        "\"Confidential Information\" means information disclosed by one party to the other in connection with the Purpose that is identified as confidential at the time of disclosure or that a reasonable person would understand to be confidential given its nature.",
        ("h2", "2. Standard of Care"),
        "Each party shall protect the other's Confidential Information with at least the same degree of care it uses to protect its own confidential information, and in any event with no less than reasonable care.",
        ("h2", "3. Permitted Disclosures"),
        "A party may disclose Confidential Information to (a) its personnel, advisers, contractors, and affiliates with a need to know for the Purpose; (b) regulators or courts where required by law; (c) any third party with the other party's prior written consent.",
        ("h2", "4. Exclusions"),
        "Confidential Information does not include information that (a) is or becomes publicly available without breach; (b) was known to the receiving party prior to disclosure; (c) is independently developed; (d) is received from a third party free of obligation; (e) is required to be disclosed by law (with prompt notice to the disclosing party).",
        ("h2", "5. Term"),
        "This Agreement commences on signature and continues for three (3) years. Obligations of confidence survive for a further three (3) years from disclosure.",
        ("h2", "6. Personal Data"),
        "Where the disclosure includes personal data, the parties shall comply with the UK GDPR and the EU GDPR (as applicable) and shall enter into a separate data-processing agreement before any such disclosure.",
        ("h2", "7. Return"),
        "On request, each party shall return or destroy the other's Confidential Information, except for back-up copies retained in the ordinary course or copies retained to demonstrate compliance with this Agreement.",
        *std_governing("the laws of England and Wales"),
        *std_misc(),
    ],
)


# --------------------------------------------------------------------
# 9. IRIS CONSULTING — mutual; long survival period (YELLOW)
# --------------------------------------------------------------------
render(
    "nda-iris-consulting.docx",
    [
        ("h1", "MUTUAL NON-DISCLOSURE AGREEMENT"),
        *std_recitals(
            f"{OUR_NAME}",
            "Iris Consulting Group Limited",
            "Iris's potential engagement to advise Stanford on supply-chain ESG performance"
        ),
        ("h2", "1. Definition"),
        "\"Confidential Information\" means any information of a confidential nature disclosed by one party to the other in connection with the Purpose, whether or not marked.",
        ("h2", "2. Obligations"),
        "Each party shall use the other's Confidential Information only for the Purpose, limit access to its personnel and professional advisers with a need to know, and protect it with reasonable care.",
        ("h2", "3. Exclusions"),
        "Standard carve-outs apply (publicly known, prior knowledge, third-party receipt, independent development, compelled disclosure).",
        ("h2", "4. Term and Survival"),
        "This Agreement is in force for three (3) years. The receiving party's obligations of confidence shall survive **for a further ten (10) years** from the date of disclosure of any specific item of Confidential Information.",
        ("h2", "5. Non-Solicitation"),
        "Each party agrees not to solicit any employee of the other party with whom it has had contact during the Purpose, for a period of twelve (12) months from termination.",
        ("h2", "6. Return"),
        "On request, each party shall return or destroy the other's Confidential Information.",
        *std_governing("the laws of England and Wales"),
        *std_misc(),
    ],
)


# --------------------------------------------------------------------
# 10. JADESTONE MFG — bilateral, balanced (GREEN)
# --------------------------------------------------------------------
render(
    "nda-jadestone-mfg.docx",
    [
        ("h1", "CONFIDENTIALITY AGREEMENT"),
        *std_recitals(
            "Jadestone Manufacturing Ltd (registered in England and Wales no. 09876543)",
            f"{OUR_NAME}",
            "Jadestone's evaluation of Stanford as a long-term supplier of precision fasteners"
        ),
        ("h2", "1. Definitions"),
        "\"Confidential Information\" means information disclosed by Jadestone to Stanford in connection with the Purpose that is marked confidential or identified as such at the time of disclosure, or that a reasonable person would understand to be confidential given its nature.",
        ("h2", "2. Use"),
        "Stanford shall use the Confidential Information solely for the Purpose and shall not disclose it other than to personnel and advisers with a need to know who are bound by equivalent obligations of confidence.",
        ("h2", "3. Carve-outs"),
        "Standard carve-outs apply (publicly known, prior knowledge, third-party receipt, independent development, compelled disclosure with prompt notice).",
        ("h2", "4. Term"),
        "This Agreement commences on signature and continues for two (2) years. Obligations of confidence survive for a further three (3) years after termination.",
        ("h2", "5. Residuals"),
        "Stanford shall not be in breach of this Agreement by virtue of any use of residual information retained in the unaided memory of its personnel.",
        ("h2", "6. Equitable Relief"),
        "Each party acknowledges that damages alone may not be sufficient remedy for breach.",
        *std_governing("the laws of England and Wales"),
        *std_misc(),
    ],
)


if __name__ == "__main__":
    files = sorted(OUT.glob("nda-*.docx"))
    print(f"Generated {len(files)} NDA fixtures:")
    for f in files:
        print(f"  {f.name}  ({f.stat().st_size} bytes)")
